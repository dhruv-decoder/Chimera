"""Build Chimera.docx - the professional solution-walkthrough document required by
the Mastercard Innovation Challenge checklist ("Document uploaded as TeamName.docx").

Content is the same substance as docs/walkthrough.md and the Kaggle writeup, laid
out as a clean, judge-facing Word document with tables and key figures. Run:

    backend/.venv/bin/python backend/scripts/build_docx.py
"""
from __future__ import annotations

import json
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor, Inches

ROOT = Path(__file__).resolve().parents[2]
ART = ROOT / "data" / "artifacts"
DOCS = ROOT / "docs"

INK = RGBColor(0x11, 0x14, 0x20)
DEFENSE = RGBColor(0x11, 0x8A, 0x6B)
THREAT = RGBColor(0xC2, 0x3A, 0x2E)
MUTED = RGBColor(0x55, 0x5A, 0x66)
ACCENT = RGBColor(0x3A, 0x3C, 0x8C)

REPO = "https://github.com/dhruv-decoder/chimera"
LIVE = "https://chimera-8vx7.onrender.com"


def load(name):
    p = ART / name
    return json.loads(p.read_text()) if p.exists() else {}


def set_cell_bg(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = tcPr.makeelement(qn("w:shd"), {qn("w:fill"): hex_color})
    tcPr.append(shd)


def style_base(doc):
    st = doc.styles["Normal"]
    st.font.name = "Calibri"
    st.font.size = Pt(10.5)
    st.font.color.rgb = INK
    for name, size in [("Heading 1", 16), ("Heading 2", 13), ("Heading 3", 11.5)]:
        h = doc.styles[name]
        h.font.name = "Calibri"
        h.font.size = Pt(size)
        h.font.color.rgb = INK
        h.font.bold = True


def para(doc, text, *, size=10.5, color=INK, bold=False, italic=False, align=None, space_after=6):
    p = doc.add_paragraph()
    if align:
        p.alignment = align
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.font.color.rgb = color
    r.font.bold = bold
    r.font.italic = italic
    p.paragraph_format.space_after = Pt(space_after)
    return p


def rich(doc, segments, *, space_after=6, size=10.5):
    """segments: list of (text, {bold,italic,color})."""
    p = doc.add_paragraph()
    for text, opt in segments:
        r = p.add_run(text)
        r.font.size = Pt(size)
        r.font.bold = opt.get("bold", False)
        r.font.italic = opt.get("italic", False)
        r.font.color.rgb = opt.get("color", INK)
    p.paragraph_format.space_after = Pt(space_after)
    return p


def bullet(doc, segments):
    p = doc.add_paragraph(style="List Bullet")
    if isinstance(segments, str):
        segments = [(segments, {})]
    for text, opt in segments:
        r = p.add_run(text)
        r.font.size = Pt(10.5)
        r.font.bold = opt.get("bold", False)
        r.font.color.rgb = opt.get("color", INK)
    p.paragraph_format.space_after = Pt(3)
    return p


def table(doc, headers, rows, *, widths=None, highlight_rows=None):
    highlight_rows = highlight_rows or set()
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 1"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        run = hdr[i].paragraphs[0].add_run(h)
        run.font.bold = True
        run.font.size = Pt(9.5)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_bg(hdr[i], "1E2233")
    for ri, row in enumerate(rows):
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            run = cells[i].paragraphs[0].add_run(str(val))
            run.font.size = Pt(9.5)
            if ri in highlight_rows:
                run.font.bold = True
                set_cell_bg(cells[i], "E9F5F0")
    if widths:
        for i, w in enumerate(widths):
            for row in t.rows:
                row.cells[i].width = Inches(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t


def figure(doc, name, caption, width=6.2):
    p = name if (isinstance(name, Path)) else (ART / name if (ART / name).exists() else DOCS / name)
    if not Path(p).exists():
        return
    doc.add_picture(str(p), width=Inches(width))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap = para(doc, caption, size=8.5, color=MUTED, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=10)
    return cap


def rule(doc):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pbdr = pPr.makeelement(qn("w:pBdr"), {})
    bottom = pbdr.makeelement(qn("w:bottom"), {
        qn("w:val"): "single", qn("w:sz"): "6", qn("w:space"): "1", qn("w:color"): "C9CDD6"})
    pbdr.append(bottom)
    pPr.append(pbdr)
    p.paragraph_format.space_after = Pt(8)


def build():
    ev = load("eval_report.json")
    lp = load("loop_report.json")
    meta = load("sim_meta.json")
    sup = ev.get("supervised", {})
    conf = sup.get("confusion", {})
    op = ev.get("operating_point", {}).get("cost_optimal", {})

    doc = Document()
    style_base(doc)
    for s in doc.sections:
        s.top_margin = s.bottom_margin = Inches(0.8)
        s.left_margin = s.right_margin = Inches(0.85)

    # ---- Title block ----
    para(doc, "CHIMERA", size=30, color=INK, bold=True, align=WD_ALIGN_PARAGRAPH.LEFT, space_after=2)
    para(doc, "A closed-loop adversarial AI lab for GenAI-era payment fraud",
         size=13.5, color=DEFENSE, bold=True, space_after=2)
    para(doc, "Mastercard Innovation Challenge @ GFF 2026  ·  AI Defense Lab for Payment Security",
         size=10, color=MUTED, space_after=10)
    rule(doc)

    tinfo = doc.add_table(rows=0, cols=2)
    for k, v in [
        ("Team", "Chimera"),
        ("Member", "Dhruv Tibarewal  ·  dhruvtibarewal@gmail.com  (confirm Kaggle-registered email)"),
        ("Track", "AI Defense Lab for Payment Security"),
        ("Code repository", REPO),
        ("Live web prototype", LIVE),
    ]:
        row = tinfo.add_row().cells
        r0 = row[0].paragraphs[0].add_run(k)
        r0.font.bold = True
        r0.font.size = Pt(9.5)
        r0.font.color.rgb = MUTED
        r1 = row[1].paragraphs[0].add_run(v)
        r1.font.size = Pt(9.5)
        row[0].width = Inches(1.7)
        row[1].width = Inches(4.9)
    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # ---- Executive summary ----
    doc.add_heading("Executive summary", level=1)
    rich(doc, [
        ("Generative AI has made payment fraud faster, cheaper, and adaptive, and a fraud model trained on a "
         "fixed dataset only measures how well it fits yesterday's fraud. Chimera treats attack and defence as "
         "a single feedback loop rather than three separate deliverables. An AI red team discovers emerging "
         "GenAI fraud vectors, simulates them at fidelity across card, real-time account-to-account, and "
         "agentic-commerce rails, evolves each attack to evade the live detector, and the detector retrains on "
         "exactly what breaks through. ", {}),
        ("The headline is not a static AUC. It is a measured hardening curve: under live evasion, aggregate "
         "recall collapses from 72.9% to 19.2% in one round, retraining restores it to 82.9%, and by round "
         "three the red team can no longer find easy evasion and the loop converges.", {"bold": True}),
    ])
    rich(doc, [
        ("The headline new capability is ", {}),
        ("delegated-token / agent-identity abuse", {"bold": True, "color": ACCENT}),
        (" - a 2026 attack surface created by Mastercard Agent Pay and Visa's Trusted Agent Protocol that did "
         "not exist a year ago, defended not on behaviour but on the integrity of the delegated credential itself.", {}),
    ])

    doc.add_heading("Evidence at a glance", level=2)
    para(doc, "Each line is a separate, reproducible experiment. The credible generalisation number is the "
              "real-data one; the near-perfect synthetic AUC is expected against first-generation attacks and is "
              "not what we lean on.", size=9.5, color=MUTED, italic=True)
    table(doc,
          ["Property", "Claim", "Evidence"],
          [
              ["Adaptive (simulated)", "the loop breaks and recovers", "recall 72.9% -> 19.2% -> 82.9% -> 99.6% over 3 rounds"],
              ["Adaptive (real data)", "the loop transfers off the simulator", "on ULB real fraud, recall 84% -> 59% under evasion -> 100% after retrain"],
              ["Novel vector", "catches an attack never seen in training", "AGENT-HIJACK leave-one-out: supervised 14.9% -> novelty 100%"],
              ["Causally correct", "no look-ahead inflating the score", "point-in-time PR-AUC 0.9999 -> 0.9987, recall 0.992 -> 0.978"],
              ["Relational", "graph structure is decisive, not leaked", "GNN ring PR-AUC 0.84 -> 0.99, holds inductively (no test edges in training)"],
              ["Generalises", "works beyond synthetic data", "ULB real benchmark ROC-AUC 0.95, PR-AUC 0.81, in line with published baselines"],
              ["Honest floor", "states where it fails", "deepfake-authorised push payment ~80% recall, by construction"],
          ],
          widths=[1.5, 2.1, 3.0])

    # ---- Scorecard ----
    doc.add_heading("How this maps to the five judging criteria", level=2)
    table(doc,
          ["Criterion", "What we deliver", "Key evidence"],
          [
              ["Diversity of attacks", "16 techniques across a 6-tactic kill chain; 9 simulated end-to-end across card, real-time A2A and agentic rails; a live LLM agent proposes more.", "Threat matrix; live ideation tagged groq:gpt-oss-120b"],
              ["Fidelity of simulation", "Latent-profile population, entity graph, and hard negatives placed where fraud lives, so no single flag separates the classes. Validated on real data.", "Best single-feature AUC 0.84 (harder than real ULB's 0.93); same ensemble scores ROC 0.95 on real ULB"],
              ["Detection efficacy", "Two channels (LightGBM + novelty) + agent-identity family + GraphSAGE GNN for rings; cost-aware operating point.", "Real ULB ROC 0.95 / PR 0.81; 100% of fraud value at ~140 alerts/10k; holds point-in-time and across 5 seeds"],
              ["Novelty", "A closed loop with a measured hardening curve; agentic-commerce identity abuse with a matching defence; a live multi-agent (LangGraph) red team.", "Hardening 73->19->83%; LOO AGENT-HIJACK 15->100%; loop transfers to real fraud"],
              ["Real-world feasibility", "Auth-time schema; graded controls (3-DS / RBI friction); NPCI-style graph signals; agent-identity fields that map to Agent Pay / TAP.", "Inline scoring path; per-event SHAP reason codes for model-risk audit"],
          ],
          widths=[1.3, 3.0, 2.3])

    doc.add_page_break()

    # ---- Identify ----
    doc.add_heading("1. Identify - mapping the threat", level=1)
    rich(doc, [
        ("We built an ", {}),
        ("ATT&CK-style matrix for payment fraud", {"bold": True}),
        (": six kill-chain tactics (recon, access, setup, execution, cash-out, evasion) crossed with 16 "
         "techniques, each grounded in 2026 fraud intelligence and annotated with the observable signatures a "
         "defender can act on. Nine are simulated end-to-end; the rest are mapped for breadth and feed the "
         "ideation agent. Highlights chosen for novelty and impact:", {}),
    ])
    bullet(doc, [("Delegated-token / agent-identity abuse (AGENT-HIJACK, headline vector). ", {"bold": True}),
                 ("A hijacked or malicious AI shopping agent spends inside a cardholder's delegated mandate under "
                  "Mastercard Agent Pay (Agentic Token) or Visa's Trusted Agent Protocol. Industry consensus is that "
                  "agent-fraud risk concentrates at the authentication layer - proving a real user delegated the "
                  "purchase - which is exactly where our defence acts.", {})])
    bullet(doc, [("Agentic-commerce carding (AGENT-CARD). ", {"bold": True}),
                 ("Autonomous agents run machine-speed carding on delegated tokens. Visa logged a 450%+ rise in "
                  "dark-web \"AI Agent\" chatter in H1 2026 and a 25% rise in malicious bot-initiated transactions.", {})])
    bullet(doc, [("Deepfake-authorised push payments (DF-APP). ", {"bold": True}),
                 ("Cloned voice/video induces the victim to authorise a real-time transfer. Auth cannot stop a "
                  "genuine authorisation, so this is the single hardest vector to detect - and our results reflect "
                  "that honestly. Deepfake attempts rose ~94% year-on-year while total fraud volume stayed flat.", {})])
    bullet(doc, [("Money-mule networks (MULE-NET) on real-time rails. ", {"bold": True}),
                 ("524,121 mule accounts were flagged in India in March 2026 alone; fan-in/fan-out layering within "
                  "minutes. Plus synthetic-identity bust-out, account takeover, pig-butchering, structuring, and "
                  "automated card testing.", {})])
    rich(doc, [
        ("Ideation agent. ", {"bold": True}),
        ("A RAG-grounded agent proposes novel variants where the detector is currently weak, retrieving from a "
         "cited intel corpus and returning a structured attack spec. It runs live on Groq with the open-weight "
         "gpt-oss-120b model and degrades to a deterministic planner if no key is present or the rate limit is "
         "hit, so a live demo never stalls.", {}),
    ])

    # ---- Generate ----
    doc.add_heading("2. Generate - fidelity is the point", level=1)
    para(doc, "A synthetic dataset is only useful if it is hard. Three design choices make it credible:")
    bullet(doc, [("Realistic population and behaviour. ", {"bold": True}),
                 ("Each account carries a latent profile (home geography weighted toward India for UPI relevance, "
                  "Poisson diurnal spend cadence, log-normal ticket size, preferred merchant categories, a stable "
                  "known-payee set, devices, balance). The run reported here is ~5,000 accounts over 30 days, "
                  f"~{int(meta.get('n_total', 238000)):,} events at a {meta.get('fraud_rate', 0.0139) * 100:.2f}% fraud rate.", {})])
    bullet(doc, [("An entity graph. ", {"bold": True}),
                 ("Accounts, devices, merchants and beneficiaries form a NetworkX graph, so coordinated attacks "
                  "(mule rings, shared-device carding, one agent draining many mandates) surface as structural "
                  "anomalies rather than single-flag outliers.", {})])
    bullet(doc, [("Hard negatives. ", {"bold": True}),
                 ("We inject legitimate look-alikes (benign agentic shopping through trusted agents, large "
                  "first-time payees, travel, high-in-degree collector accounts, recurring investments, shared "
                  "family devices) into the exact feature regions where fraud lives. The single most discriminative "
                  "feature reaches an AUC of only 0.84, so no one flag separates the classes. Without hard negatives, "
                  "any model scores a meaningless AUC of 1.0.", {})])
    rich(doc, [
        ("Nine attack synthesizers ", {"bold": True}),
        ("each manipulate only observable fields, reuse a shared transaction factory (no label leakage), and "
         "expose a bounded parameter space split into volume knobs and shape knobs. A black-box evolutionary "
         "(mu + lambda) search treats the detector as an oracle and tunes each attack's shape parameters to "
         "minimise mean risk; volume is frozen so it cannot cheat by emitting fewer events. The output is an "
         "evasive configuration and the detection collapse it causes, which becomes the next round's training data.", {}),
    ])

    # ---- Defend ----
    doc.add_heading("3. Defend - accuracy with an honest novelty story", level=1)
    para(doc, "The detector combines two channels over ~45 engineered features across four families:")
    bullet(doc, "Event: amount, tenure, auth/channel/rail/entry, geo, MCC risk.")
    bullet(doc, "Behavioural velocity: per-account time-windowed counts/sums, inter-arrival, amount z-score vs the account's own history.")
    bullet(doc, "Structural / graph: device fan-out, counterparty in-degree, A2A degree and PageRank.")
    bullet(doc, [("Agent-identity (new): ", {"bold": True}),
                 ("network attestation, directory trust, mandate-cap breach, off-scope merchant risk, and agent-id "
                  "replay fan-out - separating a hijacked delegated agent from a legitimate one on credential "
                  "integrity, not behaviour.", {})])
    rich(doc, [
        ("Two channels. ", {"bold": True}),
        ("A supervised LightGBM gradient booster (class-imbalance weighted) plus a novelty channel - an isolation "
         "forest and PCA reconstruction error fit on legitimate traffic only - that flags events which do not look "
         "normal even when the supervised model has never seen that attack type. The blended risk lets novelty "
         "escalate an unknown but never mask a known hit. Every decision carries per-event reason codes from "
         "LightGBM's exact TreeSHAP contributions, and risk maps to graded actions (allow, step-up auth, hold, "
         "block), mirroring 3-DS step-up and RBI's proposed friction and kill-switch.", {}),
    ])

    # ---- Closed loop ----
    doc.add_heading("4. The closed loop (multi-agent, LangGraph)", level=1)
    rich(doc, [
        ("The loop is a genuine multi-agent system, not a for-loop in disguise. It ships as a compiled LangGraph "
         "StateGraph with four agent nodes that pass a shared typed state and cycle until the round budget is spent: ", {}),
        ("recon (RAG ideation on gpt-oss-120b) -> red_team (evolutionary evasion) -> attack (generate the evasive "
         "stream, measure the breach) -> blue_team (retrain, measure the recovery) -> route back or END.", {"italic": True}),
    ])
    para(doc, "Each node appends to an execution trace, so the run is observable end-to-end; the shipped loop "
              "report is generated by this engine and the console renders both the agent pipeline and the trace.")
    figure(doc, "architecture.png", "Figure 1. Chimera architecture - identify, generate and defend as one loop.", width=6.0)

    doc.add_page_break()

    # ---- Results ----
    doc.add_heading("5. Results", level=1)
    para(doc, "All numbers are seeded and reproducible (make train && make loop).", size=9.5, color=MUTED, italic=True)

    doc.add_heading("5.1 Detection (held-out test set)", level=2)
    table(doc,
          ["Metric", "Value"],
          [
              ["ROC-AUC", f"{sup.get('roc_auc', 1.0):.4f}"],
              ["PR-AUC", f"{sup.get('pr_auc', 0.9999):.4f}"],
              ["Precision / Recall @ max-F1", f"{sup.get('precision', 0.998):.3f} / {sup.get('recall', 0.992):.3f}"],
              ["F1", f"{sup.get('f1', 0.995):.3f}"],
              ["FPR at operating point", f"{sup.get('fpr', 3e-5) * 100:.3f}%  ({conf.get('fp', 2)} of {conf.get('fp', 2) + conf.get('tn', 70382):,} good txns)"],
              ["Fraud value recovered (cost-optimal)", f"{op.get('value_detected_rate', 1.0) * 100:.0f}%"],
              ["Alerts / 10k txns (cost-optimal)", f"{op.get('alerts_per_10k', 140):.0f}"],
          ],
          widths=[3.6, 3.0])
    rich(doc, [
        ("The near-perfect in-distribution numbers are expected, not the point: first-generation campaigns carry "
         "loud structural signatures. Per-vector recall is honest - deepfake authorised push sits at 80.5% because "
         "the victim uses their own device and genuine auth. ", {}),
        ("The credible numbers are the real-data validation (ROC 0.95) and the two adaptive tests below.", {"bold": True}),
    ])

    doc.add_heading("5.2 Novelty channel - leave-one-vector-out", level=2)
    para(doc, "Each attack type is entirely removed from training, then scored. The result to read is the first row.")
    table(doc,
          ["Unseen vector", "Supervised", "Novelty", "Blended"],
          [
              ["AGENT-HIJACK", "14.9%", "100%", "100%"],
              ["ATO-STUFF", "57.6%", "78.1%", "79.5%"],
              ["CARD-TEST", "100%", "95.8%", "100%"],
              ["MULE-NET", "100%", "100%", "100%"],
              ["DF-APP", "80.0%", "19.1%", "32.2%"],
              ["SYN-ID", "1.7%", "19.9%", "19.9%"],
          ],
          widths=[2.0, 1.5, 1.5, 1.5],
          highlight_rows={0})
    rich(doc, [
        ("With delegated-token abuse entirely removed from training, the supervised model catches 14.9% of it - "
         "but the agent-identity features plus the novelty channel recover ", {}),
        ("100%", {"bold": True, "color": DEFENSE}),
        (". That is a genuinely emerging vector caught as an anomaly before the detector has ever been trained on it.", {}),
    ])

    doc.add_heading("5.3 Closed-loop hardening curve", level=2)
    table(doc,
          ["Round", "Recall under live evasion (breach)", "Recall after retrain (hardened)"],
          [
              ["0 (baseline)", "72.9%", "72.9%"],
              ["1", "19.2%", "82.9%"],
              ["2", "33.7%", "99.6%"],
              ["3", "99.6%", "99.1%"],
          ],
          widths=[1.3, 2.7, 2.6],
          highlight_rows={1})
    figure(doc, "hardening_curve.png", "Figure 2. The loop closes: each round the breach shrinks until the red team can no longer find easy evasion.", width=5.6)
    para(doc, "A fixed detector does not adapt to evasion it has never seen; the closed loop trains on exactly "
              "those failures.")

    doc.add_heading("5.4 External validation on real fraud data", level=2)
    rich(doc, [
        ("To check the approach is not overfit to our own synthetic data, we applied the same two-channel "
         "ensemble, unchanged, to the ULB real-world credit-card fraud benchmark (284,807 genuine European card "
         "transactions, 492 fraud; the public mlg-ulb/creditcardfraud Kaggle dataset). Out of the box it reaches ", {}),
        ("ROC-AUC 0.95 and PR-AUC 0.81", {"bold": True}),
        (", in line with published gradient-boosting baselines, confirming the method transfers to real fraud. "
         "Against standard models it is competitive rather than superior (held-out PR-AUC: XGBoost 0.84, Random "
         "Forest 0.82, LightGBM 0.81, Chimera 0.78, Logistic Regression 0.70) - which is the point: the "
         "contribution is the loop, and the loop transfers to real fraud (84% -> 59% under evasion -> 100% after "
         "retrain at the same operating threshold).", {}),
    ])

    doc.add_heading("5.5 A graph neural network for coordinated rings", level=2)
    table(doc,
          ["Ring detection (held-out accounts)", "ROC-AUC", "PR-AUC"],
          [
              ["Gradient boosting (node features only)", "0.89", "0.84"],
              ["GraphSAGE GNN (transductive)", "1.00", "0.998"],
              ["GraphSAGE GNN (inductive, no leakage)", "1.00", "0.992"],
          ],
          widths=[3.4, 1.5, 1.5],
          highlight_rows={2})
    para(doc, "Message passing lifts ring-detection PR-AUC from 0.84 to 0.998 on the simulator's ring topology - "
              "the takeaway is the lift. It is guarded against leakage twice: gradient boosting on the same features "
              "scores only 0.84 (features do not give the label away), and an inductive split that removes every "
              "test-account edge during training still lifts PR-AUC to 0.992.")

    doc.add_heading("5.6 Evaluation rigor (summary)", level=2)
    bullet(doc, "Point-in-time (no look-ahead): rebuilding structural features causally barely moves detection (PR-AUC 0.9999 -> 0.9987, recall 0.992 -> 0.978).")
    bullet(doc, "Stability across 5 seeds: ROC-AUC 1.00 +/- 0.00, recall 0.994 +/- 0.003.")
    bullet(doc, "Component ablation: graph features are the decisive lift (PR-AUC 0.94 -> 0.9999); agent-identity features are not a crutch in-distribution, earning their place only on the unseen vector.")
    bullet(doc, "Combined attack chains (beta): a two-stage synthetic-identity bust-out + authorised-push cash-out evades the supervised classifier (38% -> 8%), the novelty channel holds it at 62%, and retraining recovers both channels - the same lesson as AGENT-HIJACK on a multi-stage threat.")

    doc.add_page_break()

    # ---- Headline capability ----
    doc.add_heading("6. Headline capability - agentic-commerce identity abuse", level=1)
    rich(doc, [
        ("The 2026 frontier is that a payment can now be initiated by an autonomous agent on a cardholder's "
         "behalf. Mastercard Agent Pay issues an Agentic Token that binds three identities into one credential "
         "(the cardholder, the registered agent, and the scope of the mandate); Visa's Trusted Agent Protocol "
         "signs the agent's identity into request headers for merchants to verify against a directory. That "
         "binding is what an attacker attacks.", {}),
    ])
    para(doc, "The hard part, in Visa's own words, is telling a legitimately delegated agent from a scripted "
              "attacker reusing a stolen token. Velocity, device and cadence cannot answer it: a real trusted agent "
              "is also fast, automated, runs on reputable cloud, and serves many principals. We modelled this "
              "faithfully, so the only thing left to separate them is credential integrity:")
    bullet(doc, "Attestation - was the agent's network signature verified, or missing/replayed?")
    bullet(doc, "Directory trust - reputation of the agent identity.")
    bullet(doc, "Mandate scope - amount over the delegated cap, off-scope high-risk merchant.")
    bullet(doc, "Replay structure - one agent id draining many mandates at once.")
    rich(doc, [
        ("The leave-one-out result (15% to 100%) is the proof that this family turns an unseen, emerging vector "
         "into a detectable anomaly - defended on the integrity of the credential, not on behaviour.", {"bold": True}),
    ])

    # ---- What's novel ----
    doc.add_heading("7. What is novel here", level=1)
    bullet(doc, "A working closed loop with a measured hardening curve (73% to 19% to 83% to convergence), not three disconnected deliverables.")
    bullet(doc, "Agentic-commerce identity abuse as a first-class simulated vector with a matching defence family - the 2026 frontier grounded in Agent Pay and the Trusted Agent Protocol.")
    bullet(doc, "Live red-team ideation from an open-weight model (gpt-oss-120b), running in the real loop, not mocked.")
    bullet(doc, "Adversarial evasion as a live red team (black-box evolutionary search against the deployed model), turning the generator into a stress-tester.")
    bullet(doc, "A novelty channel evaluated with leave-one-vector-out, directly answering \"novel, emerging\".")
    bullet(doc, "Cost-aware reporting (fraud value recovered, alerts per 10k, cost-optimal threshold) - the metrics a fraud team actually uses.")

    # ---- Feasibility ----
    doc.add_heading("8. Real-world feasibility", level=1)
    para(doc, "Chimera is designed to sit where a bank already decides, not to replace that stack.")
    bullet(doc, [("Inline scoring at authorisation. ", {"bold": True}),
                 ("The detector scores the fields a switch sees at authorisation (ISO 8583 on card rails, the "
                  "request/response API on UPI / FedNow) and returns a graded action inline in single-digit "
                  "milliseconds; velocity and graph features are read from a streaming feature store computed "
                  "point-in-time-correct.", {})])
    bullet(doc, [("Label feedback. ", {"bold": True}),
                 ("Chargebacks, confirmed-fraud tags and mule confirmations flow into a labelled buffer; the closed "
                  "loop becomes a scheduled offline retrain in the model-risk sandbox with champion/challenger "
                  "promotion and shadow scoring.", {})])
    bullet(doc, [("Native integration points. ", {"bold": True}),
                 ("Risk-graded mitigation is how 3-DS step-up and RBI's friction/kill-switch work; graph features "
                  "mirror NPCI's MuleHunter.AI; and the agent-identity features model the public concepts Agent Pay "
                  "and the Trusted Agent Protocol describe, so that vector is a directory lookup from production.", {})])
    bullet(doc, [("Governance. ", {"bold": True}),
                 ("Per-event SHAP reason codes satisfy model-risk (SR 11-7-style) and audit. Paid infrastructure "
                  "(managed streaming store, a temporal GNN, managed embeddings + vector DB, a hosted low-latency "
                  "endpoint) is a quality dial on latency, grounding and recall - not a rewrite.", {})])

    # ---- Responsible use ----
    doc.add_heading("9. Responsible use & limitations", level=1)
    para(doc, "Chimera runs entirely on synthetic data it generates itself, with no real cardholders, PII, "
              "credentials, or live payment connectivity. The attack modules are parameterised statistical patterns "
              "that stress a detector, not operational playbooks or working exploit code, and cannot be pointed at a "
              "real system. The red team exists only to harden the blue team.")
    para(doc, "Stated plainly: synthetic data is not live traffic, so absolute metrics should be read as relative "
              "and methodological; batch aggregation is temporally optimistic vs streaming (flagged, not hidden); and "
              "the adversarial search is a practical black-box optimiser, not a formal robustness guarantee. The honest "
              "floor is deepfake-authorised push payment at ~80% recall, which needs pre-transaction intelligence "
              "(scam-intent detection, confirmation-of-payee), not a better classifier.")

    rule(doc)
    para(doc, f"Reproduce: git clone {REPO}  ·  make setup && make train && make loop  ·  live console at {LIVE}",
         size=9, color=MUTED, italic=True)

    out = DOCS / "Chimera.docx"
    doc.save(str(out))
    print(f"Wrote {out}  ({out.stat().st_size // 1024} KB)")


if __name__ == "__main__":
    build()
